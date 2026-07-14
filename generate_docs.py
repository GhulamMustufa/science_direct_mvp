from docx import Document

# Helper to generate a long body of text to pass the 1000 word count
long_paragraph = "This is a dummy sentence designed to take up some space. " * 30  # ~300 words
long_body = "\n\n".join([long_paragraph] * 4) # ~1200 words

short_body = "This is a short body of text, under 1000 words but more than 100 words. " * 15 # ~150 words

def create_doc(filename, title, has_abstract, has_keywords, has_sections, word_count_level, ref_count):
    doc = Document()
    doc.add_heading(title, level=1)
    
    if has_abstract:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph("This is the abstract for the paper. It describes the study and its results in a concise manner.")
        
    if has_keywords:
        doc.add_heading("Keywords", level=1)
        doc.add_paragraph("Science, Technology, Research, Study")
        
    body_text = long_body if word_count_level == "long" else short_body
    if word_count_level == "fail":
        body_text = "Too short."
    
    if has_sections:
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph(body_text[:len(body_text)//4])
        doc.add_heading("Methods", level=1)
        doc.add_paragraph(body_text[len(body_text)//4:len(body_text)//2])
        doc.add_heading("Results", level=1)
        doc.add_paragraph(body_text[len(body_text)//2:3*len(body_text)//4])
        doc.add_heading("Discussion", level=1)
        doc.add_paragraph(body_text[3*len(body_text)//4:])
    else:
        doc.add_heading("Main Content", level=1)
        doc.add_paragraph(body_text)
        
    doc.add_heading("References", level=1)
    for i in range(ref_count):
        doc.add_paragraph(f"[{i+1}] Author Name. Title of the paper {i}. Journal Name, 2024.")
        
    doc.save(filename)
    print(f"Created {filename}")

# Success Cases (3)
# Need: abstract, keywords, all sections, >1000 words, >10 refs
create_doc("success_case_1.docx", "Complete Quantum Computing Review", True, True, True, "long", 12)
create_doc("success_case_2.docx", "AI in Medical Diagnostics", True, True, True, "long", 15)
create_doc("success_case_3.docx", "Sustainable Materials Engineering", True, True, True, "long", 11)

# Warning Cases (3)
# Need: abstract, keywords, all sections, but <1000 words or <10 refs
create_doc("warning_case_1_short.docx", "Short Review on Microplastics", True, True, True, "short", 12)
create_doc("warning_case_2_few_refs.docx", "Analysis of Wind Patterns", True, True, True, "long", 7)
create_doc("warning_case_3_mixed.docx", "A brief note on Black Holes", True, True, True, "short", 6)

# Failure Cases (3)
# Need: Missing abstract, missing sections, or <100 words, or <5 refs
create_doc("failure_case_1_no_sections.docx", "No Sections Document", True, True, False, "long", 12)
create_doc("failure_case_2_no_refs.docx", "No References Document", True, True, True, "long", 2)
create_doc("failure_case_3_too_short.docx", "Extremely Short Document", True, True, True, "fail", 12)

